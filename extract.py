import docx
import sys

doc = docx.Document(sys.argv[1])
with open("C:\\Users\\CARLOS\\Music\\proyecto melissa\\pf\\extracted_docx.txt", "w", encoding="utf-8") as f:
    f.write("--- PARAGRAPHS ---\n")
    for p in doc.paragraphs:
        f.write(p.text + "\n")
    f.write("--- TABLES ---\n")
    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                row_text.append(cell.text.strip())
            f.write(" | ".join(row_text) + "\n")
